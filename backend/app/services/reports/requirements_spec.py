"""
ASTRA — Requirements Specification Document
=============================================
File: backend/app/services/reports/requirements_spec.py   ← NEW

Generates a formal Requirements Specification following
IEEE 830 / ISO 29148 structure:
  1. Cover page with project info and revision history
  2. Table of contents
  3. Requirements grouped by type, each with full details
  4. Appendices (glossary, acronyms)

Formats: docx (python-docx), pdf (reportlab)
"""

import io
from datetime import datetime
from collections import defaultdict

from sqlalchemy.orm import Session

from app.models import Requirement, Baseline, User
from app.services.reports.base import ReportGenerator, ReportOutput


TYPE_ORDER = [
    "functional", "performance", "interface", "security",
    "safety", "environmental", "reliability", "constraint",
    "maintainability", "derived",
]
TYPE_LABELS = {
    "functional": "Functional Requirements",
    "performance": "Performance Requirements",
    "interface": "Interface Requirements",
    "security": "Security Requirements",
    "safety": "Safety Requirements",
    "environmental": "Environmental Requirements",
    "reliability": "Reliability Requirements",
    "constraint": "Design Constraints",
    "maintainability": "Maintainability Requirements",
    "derived": "Derived Requirements",
}


class RequirementsSpecReport(ReportGenerator):
    name = "requirements-spec"
    supported_formats = ["docx", "pdf"]

    def generate(self, project_id: int, db: Session, options: dict | None = None) -> ReportOutput:
        options = options or {}
        fmt = options.get("format", "docx")
        project = self._get_project(db, project_id)
        reqs = self._get_requirements(db, project_id)

        # Group by type
        grouped: dict[str, list] = defaultdict(list)
        for r in reqs:
            t = self._enum_val(r.req_type)
            grouped[t].append(r)

        # Baselines for revision history
        baselines = db.query(Baseline).filter(
            Baseline.project_id == project_id
        ).order_by(Baseline.created_at.desc()).limit(10).all()

        owner = db.query(User).filter(User.id == project.owner_id).first()

        meta = {
            "project_code": project.code,
            "project_name": project.name,
            "total_requirements": len(reqs),
            "owner": owner.full_name if owner else "—",
            "generated_at": datetime.utcnow().isoformat(),
        }

        if fmt == "docx":
            return self._to_docx(project, grouped, baselines, owner, meta)
        return self._to_pdf(project, grouped, baselines, owner, meta)

    # ── DOCX ─────────────────────────────────────────────

    def _to_docx(self, project, grouped, baselines, owner, meta) -> ReportOutput:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # ── Cover Page ──
        for _ in range(4):
            doc.add_paragraph("")
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{project.code}\nRequirements Specification")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(30, 41, 59)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run(project.name)
        run2.font.size = Pt(16)
        run2.font.color.rgb = RGBColor(100, 116, 139)

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"\n\nPrepared by: {owner.full_name if owner else '—'}\n"
                     f"Date: {datetime.utcnow().strftime('%B %d, %Y')}\n"
                     f"Classification: UNCLASSIFIED").font.size = Pt(11)

        doc.add_page_break()

        # ── Revision History ──
        doc.add_heading("Revision History", level=1)
        if baselines:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            for i, h in enumerate(["Baseline", "Date", "Created By"]):
                table.rows[0].cells[i].text = h
            for bl in baselines:
                row = table.add_row()
                row.cells[0].text = bl.name
                row.cells[1].text = bl.created_at.strftime("%Y-%m-%d") if bl.created_at else "—"
                creator = None
                if bl.created_by_id:
                    from app.models import User as UserModel
                    # creator already loaded via relationship
                    row.cells[2].text = bl.created_by.full_name if bl.created_by else "—"
        else:
            doc.add_paragraph("No baselines recorded yet.")

        doc.add_page_break()

        # ── Table of Contents placeholder ──
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph(
            "[This section will be auto-populated by Word. "
            "Right-click here and select 'Update Field' after opening.]"
        )
        doc.add_page_break()

        # ── 1. Introduction ──
        doc.add_heading("1. Introduction", level=1)
        doc.add_heading("1.1 Purpose", level=2)
        doc.add_paragraph(
            f"This document specifies the requirements for the {project.name} "
            f"project ({project.code}). It follows the IEEE 830 / ISO/IEC/IEEE 29148 "
            f"standard structure for requirements specification documents."
        )
        doc.add_heading("1.2 Scope", level=2)
        doc.add_paragraph(project.description or "No project description provided.")
        doc.add_heading("1.3 Definitions & Acronyms", level=2)
        doc.add_paragraph("See Appendix A.")

        # ── 2. Overall Description ──
        doc.add_heading("2. Overall Description", level=1)
        doc.add_paragraph(
            f"The system comprises {meta['total_requirements']} requirements "
            f"across {len([t for t in grouped if grouped[t]])} categories."
        )

        # ── 3. Specific Requirements ──
        doc.add_heading("3. Specific Requirements", level=1)
        section_num = 1
        for req_type in TYPE_ORDER:
            reqs_of_type = grouped.get(req_type, [])
            if not reqs_of_type:
                continue
            label = TYPE_LABELS.get(req_type, req_type.title())
            doc.add_heading(f"3.{section_num} {label}", level=2)

            for req in reqs_of_type:
                doc.add_heading(f"{req.req_id} — {req.title}", level=3)
                # Statement
                p = doc.add_paragraph()
                p.add_run("Statement: ").bold = True
                p.add_run(req.statement)
                # Rationale
                if req.rationale:
                    p2 = doc.add_paragraph()
                    p2.add_run("Rationale: ").bold = True
                    p2.add_run(req.rationale)
                # Metadata table
                tbl = doc.add_table(rows=1, cols=5)
                tbl.style = "Light Shading Accent 1"
                for i, h in enumerate(["Priority", "Status", "Level", "Version", "Quality"]):
                    tbl.rows[0].cells[i].text = h
                row = tbl.add_row()
                row.cells[0].text = self._enum_val(req.priority)
                row.cells[1].text = self._enum_val(req.status)
                row.cells[2].text = self._enum_val(req.level) if hasattr(req, "level") else "L1"
                row.cells[3].text = str(req.version or 1)
                row.cells[4].text = str(req.quality_score or 0)
                doc.add_paragraph("")  # spacer

            section_num += 1

        # ── Appendix ──
        doc.add_page_break()
        doc.add_heading("Appendix A: Definitions & Acronyms", level=1)
        defs = [
            ("SHALL", "Mandatory requirement (binding)"),
            ("RTM", "Requirements Traceability Matrix"),
            ("CCB", "Configuration Control Board"),
            ("TBD", "To Be Determined"),
            ("TBR", "To Be Resolved"),
        ]
        for term, defn in defs:
            p = doc.add_paragraph()
            p.add_run(f"{term}: ").bold = True
            p.add_run(defn)

        buf = io.BytesIO()
        doc.save(buf)
        return ReportOutput(
            content=buf.getvalue(),
            filename=f"SRS_{project.code}_{datetime.utcnow().strftime('%Y%m%d')}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata=meta,
        )

    # ── PDF (simplified) ─────────────────────────────────

    def _to_pdf(self, project, grouped, baselines, owner, meta) -> ReportOutput:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=LETTER)
        styles = getSampleStyleSheet()
        elements = []

        elements.append(Spacer(1, 2 * inch))
        elements.append(Paragraph(
            f"<b>{project.code}</b><br/>Requirements Specification",
            ParagraphStyle("CoverTitle", parent=styles["Title"], fontSize=28, alignment=1)))
        elements.append(Spacer(1, 24))
        elements.append(Paragraph(project.name, ParagraphStyle(
            "CoverSub", parent=styles["Normal"], fontSize=16, alignment=1)))
        elements.append(Spacer(1, 48))
        elements.append(Paragraph(
            f"Prepared by: {owner.full_name if owner else '—'}<br/>"
            f"Date: {datetime.utcnow().strftime('%B %d, %Y')}",
            ParagraphStyle("CoverInfo", parent=styles["Normal"], alignment=1)))
        elements.append(PageBreak())

        # Requirements by type
        section = 1
        for req_type in TYPE_ORDER:
            reqs_of_type = grouped.get(req_type, [])
            if not reqs_of_type:
                continue
            label = TYPE_LABELS.get(req_type, req_type.title())
            elements.append(Paragraph(f"3.{section} {label}", styles["Heading1"]))
            for req in reqs_of_type:
                elements.append(Paragraph(f"<b>{req.req_id}</b> — {req.title}", styles["Heading3"]))
                elements.append(Paragraph(req.statement, styles["Normal"]))
                if req.rationale:
                    elements.append(Paragraph(f"<i>Rationale:</i> {req.rationale}", styles["Normal"]))
                elements.append(Spacer(1, 8))
            section += 1

        doc.build(elements)
        return ReportOutput(
            content=buf.getvalue(),
            filename=f"SRS_{project.code}_{datetime.utcnow().strftime('%Y%m%d')}.pdf",
            content_type="application/pdf",
            metadata=meta,
        )

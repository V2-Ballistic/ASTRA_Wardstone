"""
ASTRA — Supplier Document Pre-Extractor (Phase 7, ASTRA-TDD-INTF-002)
======================================================================
File: backend/app/services/catalog/document_extractor.py   ← NEW

The "pre-extraction" pass that turns a freshly-uploaded supplier document
(PDF / DOCX / XLSX) into a normalised intermediate representation suitable
for feeding into the LLM prompt in :mod:`app.services.catalog.icd_extractor`.

Why pre-extraction?
-------------------
Sending a 50-page binary PDF to the LLM would be wasteful and unreliable;
LLMs need clean, page-cited text plus optional table captures and (eventually)
page images. We pre-extract once locally, then craft a single tight prompt.

Outputs
-------
``ExtractedDocument`` carries:
  * ``document_type`` — "pdf" | "docx" | "xlsx" (drives prompt routing)
  * ``page_count`` — total pages (DOCX/XLSX use a synthetic pagination)
  * ``pages`` — list of :class:`ExtractedPage` (text + tables + optional image)
  * ``metadata`` — title/author/etc. (best-effort)
  * ``warnings`` — non-fatal extraction issues for the audit log

Truncation policy
-----------------
PDFs are capped at ``max_pages`` (default 50). Truncation surfaces both a
warning string AND a ``truncated=True`` flag so downstream callers can flag
in the UI ("only first 50 pages of 211 were extracted — large datasheets
need manual review").

Failure modes
-------------
  * Unsupported MIME → ``NotImplementedError``  (caller marks doc=FAILED)
  * Corrupt file       → ``DocumentExtractionError`` (caller marks doc=FAILED)
  * Camelot crash      → caught + logged as a warning; page is returned with
                          empty ``tables=[]`` so the LLM still sees the text.
"""

from __future__ import annotations

import logging
import os
from dataclasses import dataclass, field
from typing import List, Optional

logger = logging.getLogger("astra.catalog.document_extractor")


# ──────────────────────────────────────────────────────────────
#  Public dataclasses
# ──────────────────────────────────────────────────────────────

@dataclass
class ExtractedPage:
    """One page of a pre-extracted document."""
    page_number: int                                 # 1-indexed
    text: str                                        # full visible text
    image_bytes: Optional[bytes] = None              # 200 DPI render (PDF only)
    tables: List[List[List[str]]] = field(default_factory=list)
    """Camelot-style: list of tables; each table is a list of rows; each row
    is a list of cell strings. Empty list when no tables are detected."""


@dataclass
class ExtractedDocument:
    """Normalised intermediate representation, prompt-ready."""
    document_type: str                               # "pdf" | "docx" | "xlsx"
    page_count: int
    pages: List[ExtractedPage] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)
    truncated: bool = False                          # True if >max_pages

    def total_text_length(self) -> int:
        return sum(len(p.text) for p in self.pages)


class DocumentExtractionError(RuntimeError):
    """Raised on un-recoverable extraction failures (corrupt file, etc.)."""


# ──────────────────────────────────────────────────────────────
#  MIME → handler routing
# ──────────────────────────────────────────────────────────────

# PDFs
_PDF_MIMES = {"application/pdf", "application/x-pdf", "application/acrobat"}
# DOCX (Office Open XML)
_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
# XLSX (Office Open XML spreadsheet)
_XLSX_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}


def _resolve_doc_type(file_path: str, mime_type: str) -> str:
    """Pick the handler key from MIME, falling back to extension if MIME is
    generic (``application/octet-stream`` from some browsers/uploads)."""
    mime = (mime_type or "").lower()
    if mime in _PDF_MIMES:
        return "pdf"
    if mime in _DOCX_MIMES:
        return "docx"
    if mime in _XLSX_MIMES:
        return "xlsx"
    # Fall back to extension
    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    if ext == "pdf":
        return "pdf"
    if ext in {"docx", "docm"}:
        return "docx"
    if ext in {"xlsx", "xlsm", "xls"}:
        return "xlsx"
    return "unknown"


# ──────────────────────────────────────────────────────────────
#  PDF handler — PyMuPDF + camelot
# ──────────────────────────────────────────────────────────────

# Cap on per-page image bytes captured in the dataclass, to keep the
# in-memory representation small. Caller may still re-render pages on demand.
_IMAGE_PAGE_CAP = 5    # only first 5 pages get an image render
_RENDER_DPI = 200


def _extract_pdf(file_path: str, max_pages: int) -> ExtractedDocument:
    """PDF → ExtractedDocument via PyMuPDF (text + images) + camelot (tables)."""
    import fitz  # PyMuPDF

    warnings: List[str] = []
    pages: List[ExtractedPage] = []

    try:
        doc = fitz.open(file_path)
    except Exception as exc:  # pragma: no cover - corrupt-file path
        raise DocumentExtractionError(f"PyMuPDF failed to open PDF: {exc}") from exc

    try:
        total_pages = doc.page_count
        truncated = total_pages > max_pages
        if truncated:
            warnings.append(
                f"Document has {total_pages} pages; truncated to first {max_pages} for extraction"
            )

        meta = dict(doc.metadata or {})
        # Strip null-ish entries the LLM doesn't need.
        meta = {k: v for k, v in meta.items() if v}

        # Per-page text + optional image render.
        page_limit = min(total_pages, max_pages)
        zoom = _RENDER_DPI / 72.0
        matrix = fitz.Matrix(zoom, zoom)

        for pno in range(page_limit):
            page = doc.load_page(pno)
            text = page.get_text("text") or ""
            image_bytes: Optional[bytes] = None
            if pno < _IMAGE_PAGE_CAP:
                try:
                    pix = page.get_pixmap(matrix=matrix, alpha=False)
                    image_bytes = pix.tobytes("png")
                except Exception as exc:    # pragma: no cover
                    warnings.append(f"Page {pno + 1} image render failed: {exc}")
            pages.append(ExtractedPage(
                page_number=pno + 1,
                text=text,
                image_bytes=image_bytes,
                tables=[],   # filled below
            ))
    finally:
        doc.close()

    # Tables via camelot — fail-soft per page (camelot can crash on weird PDFs).
    try:
        import camelot
        # Camelot pages param is 1-indexed strings. Use lattice when possible
        # (better for ruled tables) and fall back to stream silently.
        page_spec = f"1-{min(total_pages, max_pages)}"
        try:
            tables = camelot.read_pdf(file_path, pages=page_spec, flavor="lattice")
        except Exception as exc_lattice:
            logger.debug("camelot lattice failed: %s — trying stream", exc_lattice)
            try:
                tables = camelot.read_pdf(file_path, pages=page_spec, flavor="stream")
            except Exception as exc_stream:
                warnings.append(
                    f"Camelot table extraction failed (lattice + stream): {exc_stream}"
                )
                tables = []
        # Group tables by page
        per_page: dict[int, list] = {}
        for tbl in tables or []:
            page_no = int(getattr(tbl, "page", 0))
            try:
                rows = tbl.df.values.tolist()
                rows = [[str(cell) for cell in row] for row in rows]
            except Exception:    # pragma: no cover
                rows = []
            per_page.setdefault(page_no, []).append(rows)
        for p in pages:
            p.tables = per_page.get(p.page_number, [])
    except ImportError:    # pragma: no cover
        warnings.append("camelot-py not installed — tables omitted")
    except Exception as exc:    # pragma: no cover
        warnings.append(f"camelot table extraction errored: {exc}")

    return ExtractedDocument(
        document_type="pdf",
        page_count=total_pages,
        pages=pages,
        metadata=meta,
        warnings=warnings,
        truncated=truncated,
    )


# ──────────────────────────────────────────────────────────────
#  DOCX handler — python-docx
# ──────────────────────────────────────────────────────────────

# DOCX is reflowable. We synthesise pagination by chunking paragraphs.
_DOCX_PAGE_PARAGRAPHS = 30


def _extract_docx(file_path: str, max_pages: int) -> ExtractedDocument:
    """DOCX → ExtractedDocument via python-docx. No page images (reflowable)."""
    import docx

    warnings: List[str] = []
    try:
        document = docx.Document(file_path)
    except Exception as exc:    # pragma: no cover
        raise DocumentExtractionError(f"python-docx failed to open file: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Tables (one global list captured on a virtual "page 1" for now)
    docx_tables: List[List[List[str]]] = []
    for tbl in document.tables:
        rows: List[List[str]] = []
        for row in tbl.rows:
            rows.append([cell.text for cell in row.cells])
        if rows:
            docx_tables.append(rows)

    # Synthetic pagination — group paragraphs into "pages" of N.
    pages: List[ExtractedPage] = []
    page_no = 0
    for i in range(0, len(paragraphs), _DOCX_PAGE_PARAGRAPHS):
        page_no += 1
        if page_no > max_pages:
            warnings.append(
                f"DOCX has {len(paragraphs)} paragraphs; truncated to {max_pages} synthetic pages"
            )
            break
        chunk = "\n".join(paragraphs[i: i + _DOCX_PAGE_PARAGRAPHS])
        pages.append(ExtractedPage(
            page_number=page_no,
            text=chunk,
            image_bytes=None,
            tables=docx_tables if page_no == 1 else [],
        ))

    if not pages:
        # Empty doc — still return one synthetic page with the table dump
        pages.append(ExtractedPage(page_number=1, text="", tables=docx_tables))

    meta: dict = {}
    try:
        cp = document.core_properties
        for attr in ("title", "author", "subject", "keywords", "comments"):
            v = getattr(cp, attr, None)
            if v:
                meta[attr] = str(v)
    except Exception:    # pragma: no cover
        pass

    return ExtractedDocument(
        document_type="docx",
        page_count=len(pages),
        pages=pages,
        metadata=meta,
        warnings=warnings,
        truncated=page_no > max_pages,
    )


# ──────────────────────────────────────────────────────────────
#  XLSX handler — openpyxl
# ──────────────────────────────────────────────────────────────

def _extract_xlsx(file_path: str, max_pages: int) -> ExtractedDocument:
    """XLSX → ExtractedDocument: each sheet is one synthetic page (table only)."""
    import openpyxl

    warnings: List[str] = []
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
    except Exception as exc:    # pragma: no cover
        raise DocumentExtractionError(f"openpyxl failed to open file: {exc}") from exc

    pages: List[ExtractedPage] = []
    sheet_names = wb.sheetnames
    truncated = False
    if len(sheet_names) > max_pages:
        warnings.append(
            f"Workbook has {len(sheet_names)} sheets; truncated to first {max_pages}"
        )
        sheet_names = sheet_names[:max_pages]
        truncated = True

    for idx, name in enumerate(sheet_names, start=1):
        ws = wb[name]
        rows: List[List[str]] = []
        text_chunks: List[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [("" if v is None else str(v)) for v in row]
            rows.append(cells)
            non_empty = [c for c in cells if c]
            if non_empty:
                text_chunks.append("\t".join(non_empty))
        pages.append(ExtractedPage(
            page_number=idx,
            text=f"--- Sheet: {name} ---\n" + "\n".join(text_chunks),
            tables=[rows] if rows else [],
        ))

    meta: dict = {}
    try:
        props = wb.properties
        for attr in ("title", "creator", "subject", "keywords", "description"):
            v = getattr(props, attr, None)
            if v:
                meta[attr] = str(v)
    except Exception:    # pragma: no cover
        pass

    wb.close()

    return ExtractedDocument(
        document_type="xlsx",
        page_count=len(pages),
        pages=pages,
        metadata=meta,
        warnings=warnings,
        truncated=truncated,
    )


# ──────────────────────────────────────────────────────────────
#  Public entry point
# ──────────────────────────────────────────────────────────────

def extract_document(
    file_path: str,
    mime_type: str,
    max_pages: int = 50,
) -> ExtractedDocument:
    """Pre-extract content from a supplier document.

    Parameters
    ----------
    file_path : str
        Absolute path on disk (the SupplierDocument.file_path).
    mime_type : str
        The recorded MIME type. Falls back to extension sniffing if generic.
    max_pages : int, optional
        Cap to avoid blowing up the LLM prompt on giant datasheets. Default 50.

    Returns
    -------
    ExtractedDocument

    Raises
    ------
    NotImplementedError
        Unsupported MIME / extension. Caller should mark the SupplierDocument
        ``extraction_status=FAILED`` and surface in the UI.
    DocumentExtractionError
        Recognised type but the file is corrupt / unreadable.
    """
    doc_type = _resolve_doc_type(file_path, mime_type)
    if doc_type == "pdf":
        return _extract_pdf(file_path, max_pages)
    if doc_type == "docx":
        return _extract_docx(file_path, max_pages)
    if doc_type == "xlsx":
        return _extract_xlsx(file_path, max_pages)
    raise NotImplementedError(
        f"Unsupported document type for ICD extraction: mime={mime_type!r}, "
        f"file={os.path.basename(file_path)!r}"
    )
